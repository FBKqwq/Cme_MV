from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path(
    r"C:\Users\zhurunjie\Desktop\CmePlatform\pro\data\review\state\exports"
    r"\待医师判定复验实体_2026-07-31.docx"
)

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION_FILL = "FFF8E8"
CAUTION = "7A5A00"
BORDER = "CCD5E1"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


ENTRIES = [
    {
        "index": 1,
        "name": "咽部检查",
        "document": "百日咳（2023年）",
        "document_id": "DOC_65f24d91cdee",
        "chunk_id": "DOC_65f24d91cdee_CH0004",
        "section": "正文 > 百日咳诊疗方案",
        "pages": "2-3",
        "entity_id": "OCC_7537cd53433b",
        "original_type": "病因（etiologies）",
        "proposed_result": "拟修改为检查（tests），并关联复验规范实体",
        "state": "modified + pending",
        "updated_at": "2026-07-30 09:53:36",
        "evidence": "咽部检查或遇到冷风、烟雾、进食等诱因时",
        "context": (
            "剧烈咳嗽刺激大脑皮质的咳嗽中枢可形成持久的兴奋灶，"
            "咽部检查或遇到冷风、烟雾、进食等诱因时，可引起痉挛性咳嗽发作。"
        ),
        "question": (
            "“咽部检查”在该句中是检查项目、诱发痉咳的操作性诱因，"
            "还是不应作为独立知识实体？"
        ),
        "system_note": (
            "系统已把实体类型从“病因”改为“检查”，但最终决定仍为 pending。"
            "该修改不能直接视为复验通过。"
        ),
        "options": [
            "保留为检查（tests）",
            "保留，但类型改为病因（etiologies）",
            "不作为独立实体，确认删除",
            "其他处理（请在医学依据中说明）",
        ],
    },
    {
        "index": 2,
        "name": "菌种鉴定",
        "document": "成人耳念珠菌感染诊治防控专家共识",
        "document_id": "DOC_faf109889cc7",
        "chunk_id": "DOC_faf109889cc7_CH0005",
        "section": "正文 > 3 微生物学鉴定",
        "pages": "1-2",
        "entity_id": "OCC_0b596a2dda65",
        "original_type": "检查（tests）",
        "proposed_result": "已标记删除",
        "state": "deleted + pending",
        "updated_at": "2026-07-30 09:49:01",
        "evidence": "进行菌种鉴定",
        "context": (
            "分离鉴定技术：该菌镜下形态学无特征性提示，需要菌落分纯后进行菌种鉴定。"
            "正确鉴定要求人员意识到位、经验充足、质控在控，结果有可信性、重复性、可比性。"
        ),
        "question": (
            "“菌种鉴定”是否应作为检查类知识实体保留，还是因表述过于通用而删除？"
        ),
        "system_note": (
            "系统记录为删除操作，但 review_decision 仍为 pending，"
            "删除尚未形成可交付的终态。"
        ),
        "options": [
            "保留为检查（tests）",
            "保留，但修改名称或实体类型",
            "确认删除",
            "其他处理（请在医学依据中说明）",
        ],
    },
    {
        "index": 3,
        "name": "BlauB 综合征",
        "document": "Blau综合征诊疗专家共识（2024版）",
        "document_id": "DOC_91bfb25e36f4",
        "chunk_id": "DOC_91bfb25e36f4_CH0002",
        "section": "正文 > 综合征诊疗专家共识（2024版）",
        "pages": "1",
        "entity_id": "OCC_d6e2d6a64b2d",
        "original_type": "疾病（diseases）",
        "proposed_result": "尚未修改；名称疑似存在版面/OCR污染",
        "state": "pending",
        "updated_at": "2026-07-30 11:24:40",
        "evidence": "BlauB 综合征是一种罕见的单基因突变导致的自身炎症性疾病",
        "context": (
            "【摘要】BlauB 综合征是一种罕见的单基因突变导致的自身炎症性疾病，"
            "以皮疹、肉芽肿性多关节炎和葡萄膜炎的三联征为主要临床特征。"
            "同段后文使用“Blau综合征”。"
        ),
        "question": (
            "“BlauB 综合征”是否为版面/OCR噪声，应规范为“Blau综合征”并与同名实体合并？"
        ),
        "system_note": (
            "该实体尚未产生最终复验决定。当前名称中的字母 B 与同段标准名称不一致。"
        ),
        "options": [
            "规范为“Blau综合征”，保留为疾病（diseases）",
            "保留原名称“BlauB 综合征”",
            "与已有“Blau综合征”合并，并删除重复项",
            "其他处理（请在医学依据中说明）",
        ],
    },
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    name: str = FONT_CN,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(
    paragraph,
    *,
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
    alignment=None,
    keep_with_next: bool | None = None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment
    if keep_with_next is not None:
        fmt.keep_with_next = keep_with_next


def add_text_paragraph(
    container,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    italic: bool = False,
    color: str = "222222",
    before: float = 0,
    after: float = 6,
    line_spacing: float = 1.25,
    alignment=None,
):
    paragraph = container.add_paragraph()
    set_paragraph(
        paragraph,
        before=before,
        after=after,
        line_spacing=line_spacing,
        alignment=alignment,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def format_cell_text(cell, text: str, *, bold=False, color="222222", size=10.2) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph(paragraph, before=0, after=0, line_spacing=1.2)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text_run = OxmlElement("w:r")
    text_props = OxmlElement("w:rPr")
    text_fonts = OxmlElement("w:rFonts")
    text_fonts.set(qn("w:ascii"), FONT_LATIN)
    text_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    text_props.append(text_fonts)
    text_size = OxmlElement("w:sz")
    text_size.set(qn("w:val"), "18")
    text_props.append(text_size)
    text_run.append(text_props)
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    text_run.append(text_node)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_separate)
    paragraph._p.append(text_run)
    paragraph._p.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_CN
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = FONT_CN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.clear()
    set_paragraph(
        paragraph,
        before=0,
        after=0,
        line_spacing=1.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run("CmePlatform  |  医师复核判定表")
    set_run_font(run, size=9, bold=True, color=MUTED)

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    set_paragraph(
        paragraph,
        before=0,
        after=0,
        line_spacing=1.0,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    add_page_field(paragraph)


def add_title_block(doc: Document) -> None:
    add_text_paragraph(
        doc,
        "医师复核判定表",
        size=24,
        bold=True,
        color=INK,
        after=4,
    )
    add_text_paragraph(
        doc,
        "知识复验待确认实体（批次截至 2026-07-30）",
        size=13.5,
        color=MUTED,
        after=16,
    )
    metadata = [
        ("对象", "当前复验增量中 review_decision 仍为 pending 的实体"),
        ("数量", "3 条"),
        ("编制日期", "2026-07-31"),
        ("使用方式", "勾选判定结论；必要时填写规范名称、实体类型及医学依据；签名后交回录入"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, before=0, after=3, line_spacing=1.18)
        label_run = paragraph.add_run(f"{label}：")
        set_run_font(label_run, size=10.5, bold=True, color=INK)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color="333333")


def add_one_cell_callout(doc: Document, label: str, text: str, *, fill: str, label_color: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph(paragraph, before=0, after=0, line_spacing=1.25)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, bold=True, color=label_color)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color="222222")
    add_text_paragraph(doc, "", size=2, after=2)


def add_overview(doc: Document) -> None:
    heading = doc.add_paragraph("判定说明", style="Heading 1")
    heading.paragraph_format.space_before = Pt(16)
    add_one_cell_callout(
        doc,
        "范围说明",
        (
            "本表收录全部仍未形成终态的实体：1 条修改后待确认、"
            "1 条删除后待确认、1 条普通待确认。实体判定用于知识库质量控制，"
            "不构成患者诊断或治疗建议。"
        ),
        fill=CALLOUT,
        label_color=DARK_BLUE,
    )

    table = doc.add_table(rows=1, cols=4)
    widths = [650, 1750, 2600, 4360]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["序号", "待判定实体", "来源", "核心待决问题"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        format_cell_text(cell, header, bold=True, color=INK, size=10)
        cell.paragraphs[0].alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        )
    mark_repeat_header(table.rows[0])

    summaries = [
        ("1", "咽部检查", "百日咳（2023年），第2-3页", "检查实体、诱因实体，还是不单独入库"),
        ("2", "菌种鉴定", "成人耳念珠菌感染共识，第1-2页", "保留为检查，还是确认删除"),
        ("3", "BlauB 综合征", "Blau综合征共识，第1页", "规范名称并合并，还是保留原名"),
    ]
    for row_values in summaries:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            format_cell_text(cells[i], value, size=9.5)
            if i == 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_text_paragraph(doc, "", size=2, after=2)
    add_one_cell_callout(
        doc,
        "填写要求",
        (
            "每条实体只勾选一个主结论。若选择“修改”或“其他处理”，"
            "请写明规范名称、目标实体类型及简要医学依据。"
        ),
        fill=CAUTION_FILL,
        label_color=CAUTION,
    )


def add_metadata_table(doc: Document, entry: dict) -> None:
    rows = [
        ("来源文献", entry["document"]),
        ("定位", f'{entry["section"]}；当前分块页码 {entry["pages"]}'),
        ("工程标识", f'{entry["entity_id"]}  |  {entry["chunk_id"]}'),
        ("原始实体类型", entry["original_type"]),
        ("当前拟处理", entry["proposed_result"]),
        ("复验状态", f'{entry["state"]}（最后更新：{entry["updated_at"]}）'),
    ]
    table = doc.add_table(rows=0, cols=2)
    widths = [2200, 7160]
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT_GRAY)
        format_cell_text(cells[0], label, bold=True, color=INK, size=9.7)
        format_cell_text(cells[1], value, size=9.7)
    set_table_geometry(table, widths)
    set_table_borders(table)


def add_decision_options(doc: Document, options: list[str]) -> None:
    table = doc.add_table(rows=0, cols=1)
    for option in options:
        cell = table.add_row().cells[0]
        format_cell_text(cell, f"□  {option}", size=10.5)
        cell.paragraphs[0].paragraph_format.space_before = Pt(2)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size="4")


def add_notes_box(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    set_paragraph(paragraph, after=8, line_spacing=1.2)
    run = paragraph.add_run("规范名称 / 目标类型：")
    set_run_font(run, size=10, bold=True, color=INK)
    for label in ("医学依据：", "补充意见："):
        paragraph = cell.add_paragraph()
        set_paragraph(paragraph, before=0, after=18, line_spacing=1.2)
        run = paragraph.add_run(label)
        set_run_font(run, size=10, bold=True, color=INK)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size="4")


def add_signature_table(doc: Document) -> None:
    table = doc.add_table(rows=2, cols=4)
    widths = [1450, 3150, 1450, 3310]
    labels = [
        ("判定医师", "签名："),
        ("科室 / 职称", ""),
        ("判定日期", "____年__月__日"),
        ("复核录入人", ""),
    ]
    for row_index in range(2):
        for pair_index in range(2):
            label, value = labels[row_index * 2 + pair_index]
            label_cell = table.cell(row_index, pair_index * 2)
            value_cell = table.cell(row_index, pair_index * 2 + 1)
            set_cell_shading(label_cell, LIGHT_GRAY)
            format_cell_text(label_cell, label, bold=True, color=INK, size=9.5)
            format_cell_text(value_cell, value, size=9.5)
    set_table_geometry(table, widths)
    set_table_borders(table)


def add_entry(doc: Document, entry: dict) -> None:
    heading = doc.add_paragraph(
        f'{entry["index"]}. {entry["name"]}',
        style="Heading 1",
    )
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.space_before = Pt(0)

    status = doc.add_paragraph()
    set_paragraph(status, before=0, after=10, line_spacing=1.1)
    run = status.add_run("待医师判定")
    set_run_font(run, size=10, bold=True, color=CAUTION)

    add_metadata_table(doc, entry)

    subheading = doc.add_paragraph("原文证据", style="Heading 2")
    subheading.paragraph_format.space_before = Pt(12)
    add_one_cell_callout(
        doc,
        "证据短句",
        entry["evidence"],
        fill=LIGHT_BLUE,
        label_color=DARK_BLUE,
    )
    add_one_cell_callout(
        doc,
        "上下文",
        entry["context"],
        fill=CALLOUT,
        label_color=DARK_BLUE,
    )

    subheading = doc.add_paragraph("待判定问题", style="Heading 2")
    add_one_cell_callout(
        doc,
        "医师需回答",
        entry["question"],
        fill=CAUTION_FILL,
        label_color=CAUTION,
    )
    add_text_paragraph(
        doc,
        f'系统记录说明：{entry["system_note"]}',
        size=9.6,
        color=MUTED,
        after=8,
    )

    subheading = doc.add_paragraph("判定结论", style="Heading 2")
    add_decision_options(doc, entry["options"])

    subheading = doc.add_paragraph("医学依据与修订意见", style="Heading 2")
    add_notes_box(doc)
    add_text_paragraph(doc, "", size=2, after=2)
    add_signature_table(doc)


def build_document() -> Path:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)

    props = doc.core_properties
    props.title = "医师复核判定表：知识复验待确认实体"
    props.subject = "CmePlatform 知识复验医师判定"
    props.author = "CmePlatform"
    props.keywords = "知识复验, 医师判定, 实体审核"

    add_title_block(doc)
    add_overview(doc)
    for entry in ENTRIES:
        add_entry(doc, entry)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
