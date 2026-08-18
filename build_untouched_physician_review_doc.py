from __future__ import annotations

import json
import re
from collections import defaultdict
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from build_physician_review_doc import (
    BLUE,
    BORDER,
    CALLOUT,
    CAUTION_FILL,
    CONTENT_WIDTH_DXA,
    DARK_BLUE,
    INK,
    LIGHT_GRAY,
    MUTED,
    add_text_paragraph,
    configure_page,
    configure_styles,
    format_cell_text,
    mark_repeat_header,
    set_cell_shading,
    set_paragraph,
    set_run_font,
    set_table_borders,
    set_table_geometry,
)


# 本文件位于项目根目录 Cme_MV 下。
# 因此不再写死 Windows 用户名，换电脑或换项目目录后也能正常读取数据。
PROJECT_ROOT = Path(__file__).resolve().parent
REVIEW_ROOT = PROJECT_ROOT / "data" / "review" / "1"
GENERATED_DATE = date.today().isoformat()

OUTPUT_PATH = (
    REVIEW_ROOT
    / "state"
    / "exports"
    / f"医师复验清单_机器待复验且人工未操作_上下文扩展_{GENERATED_DATE}.docx"
)

# 不再排除任何文档。当前批次中的全部文献都参与未操作实体筛选。
EXCLUDED_DOCUMENT_ID: str | None = None

# 正式 API 中待复验数量会变化，因此不再固定为 32。
# 如需临时校验固定数量，可将 None 改成具体数字。
EXPECTED_UNTOUCHED: int | None = None

# 医师复验上下文长度控制：
# - 硬性至少 200 个非空白字符；
# - 默认尽量提供 300~600 字左右的上下文；
# - 必要时自动拼接同一文献相邻 chunk，而不是只看当前 chunk。
CONTEXT_HARD_MIN = 200
CONTEXT_TARGET_MIN = 300
CONTEXT_TARGET_MAX = 600
CONTEXT_NEIGHBOR_RADIUS = 4

TYPE_LABELS = {
    "diseases": "疾病",
    "sub_diseases": "疾病亚型",
    "symptoms": "症状与体征",
    "tests": "检查",
    "treatments": "治疗原则",
    "plans": "治疗方案",
    "methods": "实施方法",
    "etiologies": "病因",
    "pathogeneses": "发病机制",
}


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def read_jsonl(path: Path):
    with path.open("r", encoding="utf-8-sig") as handle:
        for line in handle:
            line = line.strip()
            if line:
                yield json.loads(line)


def wrap_title(title: str) -> str:
    """给文献标题加上《》，如果标题本身还没有的话。"""
    if not title:
        return title
    if "《" in title and "》" in title:
        return title
    return f"《{title}》"


def sanitize_display(value: Any, *, limit: int | None = None) -> str:
    text = "" if value is None else str(value)
    pieces: list[str] = []
    for char in text:
        code = ord(char)
        if 0xE000 <= code <= 0xF8FF:
            pieces.append(f"[缺失字形U+{code:04X}]")
        elif code < 32 and char not in "\t\n\r":
            pieces.append(" ")
        else:
            pieces.append(char)
    text = "".join(pieces)
    text = re.sub(r"\s+", " ", text).strip()
    # 不再在 sanitize_display 中按 limit 硬截断，避免半句乱码。
    # 截断由 _clip_at_sentence_boundary 在句末处理。
    return text


def entity_doc_id(entity: dict[str, Any]) -> str:
    return str(entity.get("document_id") or entity.get("_doc_id") or "")


def source_chunk_id(entity: dict[str, Any]) -> str:
    doc_id = entity_doc_id(entity)
    chunk_id = str(entity.get("chunk_id") or "")
    prefix = f"{doc_id}_"
    return chunk_id[len(prefix) :] if chunk_id.startswith(prefix) else chunk_id


def load_chunks() -> tuple[
    dict[str, dict[str, Any]],
    dict[tuple[str, str], dict[str, Any]],
    list[str],
    dict[str, list[dict[str, Any]]],
]:
    """
    读取当前批次的 chunk。

    除原来的 documents / chunks / document_order 外，
    额外返回 chunk_sequences：
        doc_id -> 按原文顺序排列的 chunk 列表

    这样生成复验 context 时，如果当前 chunk 太短，可以可靠地向同一篇
    文献的前后 chunk 扩展，而不需要再从导出的 PDF 反向 OCR。
    """
    documents: dict[str, dict[str, Any]] = {}
    chunks: dict[tuple[str, str], dict[str, Any]] = {}
    document_order: list[str] = []
    chunk_sequences: dict[str, list[dict[str, Any]]] = {}

    for path in sorted((REVIEW_ROOT / "current" / "chunks").glob("*chunk.json")):
        payload = read_json(path)
        doc_id = str(payload["doc_id"])
        documents[doc_id] = {
            "document_id": doc_id,
            "title": str(payload.get("source_title") or path.stem),
            "chunk_count": len(payload.get("chunks", [])),
        }
        document_order.append(doc_id)

        ordered_chunks: list[dict[str, Any]] = []
        for index, chunk in enumerate(payload.get("chunks", []), start=1):
            item = dict(chunk)
            item["_index"] = index
            chunks[(doc_id, str(chunk["chunk_id"]))] = item
            ordered_chunks.append(item)

        chunk_sequences[doc_id] = ordered_chunks

    return documents, chunks, document_order, chunk_sequences


def load_reviewed_entity_ids() -> set[str]:
    """读取当前已经形成明确人工结论的实体 ID。

    这里判断的是实体的“当前复验结论”，而不是它是否曾经被操作过。

    因此：
    - review_decision=accepted/rejected 等终态：视为已人工复验；
    - review_decision=pending：视为当前未复验；
    - 即使 review_version > 0，或曾经修改/通过/拒绝，只要后来撤销为
      pending，就应重新进入未复验 PDF。

    同一实体若出现多条记录，优先采用 review_version 较大的最新记录。
    """
    latest_states: dict[str, tuple[tuple[int, str, str], str]] = {}
    delta_root = REVIEW_ROOT / "state" / "reviews"

    for path in sorted(delta_root.glob("*/*.review.json")):
        payload = read_json(path)
        for entity in payload.get("entities", []):
            entity_id = str(entity.get("entity_id") or "").strip()
            if not entity_id:
                continue

            try:
                version = int(entity.get("review_version", 0) or 0)
            except (TypeError, ValueError):
                version = 0

            decision = str(
                entity.get("review_decision") or "pending"
            ).strip().lower()
            updated_at = str(
                entity.get("updated_at")
                or entity.get("review_updated_at")
                or ""
            )

            rank = (version, updated_at, str(path))
            current = latest_states.get(entity_id)
            if current is None or rank >= current[0]:
                latest_states[entity_id] = (rank, decision)

    pending_decisions = {
        "",
        "pending",
        "review",
        "unreviewed",
        "not_reviewed",
    }

    return {
        entity_id
        for entity_id, (_, decision) in latest_states.items()
        if decision not in pending_decisions
    }


def load_source_entities() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    node_root = REVIEW_ROOT / "current" / "entity_nodes"
    paths = sorted(node_root.glob("*.entity_label_result.jsonl"))
    if not paths:
        paths = sorted(node_root.glob("*.entity_nodes.base.jsonl"))
    for path in paths:
        rows.extend(read_jsonl(path))
    return rows


def _visible_text_len(text: str) -> int:
    """上下文长度按“非空白字符”计算。"""
    return len(re.sub(r"\s+", "", text or ""))


def _sentence_spans(text: str) -> list[tuple[int, int]]:
    """
    将原文切成语义片段。
    不再像旧版那样“实体后遇到第一个句号就立刻截断”，
    而是允许保留前后多句。
    """
    spans: list[tuple[int, int]] = []
    start = 0

    # 分号也可以作为医学指南中的弱边界；连续换行视作段落边界。
    pattern = re.compile(r"[。！？!?；;]|\n\s*\n")

    for match in pattern.finditer(text):
        end = match.end()
        if text[start:end].strip():
            spans.append((start, end))
        start = end

    if start < len(text) and text[start:].strip():
        spans.append((start, len(text)))

    return spans


def _clip_at_sentence_boundary(text: str, max_len: int, *, min_sentences: int = 2) -> str:
    """
    将文本截断到不超过 max_len 个非空白字符，且在句末截断（而非句中）。
    保证至少保留 min_sentences 个完整句子。
    """
    stripped = text.strip()
    if _visible_text_len(stripped) <= max_len:
        return stripped

    spans = _sentence_spans(stripped)
    if len(spans) <= min_sentences:
        return stripped

    # 从后往前删句子，直到不超过 max_len，但至少保留 min_sentences 句
    kept = spans[:]
    while len(kept) > min_sentences and _visible_text_len(
        stripped[kept[0][0]:kept[-1][1]]
    ) > max_len:
        # 判断删头还是删尾，哪个能让总量更接近 max_len 且实体仍在中间
        head_len = kept[1][1] - kept[0][0]
        tail_len = kept[-1][1] - kept[-2][0]
        if head_len >= tail_len:
            kept.pop(0)
        else:
            kept.pop()

    result = stripped[kept[0][0]:kept[-1][1]].strip()
    return result


def _expand_semantic_window(
    text: str,
    *,
    mention_start: int,
    mention_end: int,
) -> str:
    """
    围绕实体所在位置扩展上下文，保证至少 2 句完整句子。
    1. 找到实体所在句；
    2. 至少再取前 1 句 + 后 1 句（共 3 句，不足时取 2 句）；
    3. 若仍不足 CONTEXT_TARGET_MIN，继续向两侧按句扩展；
    4. 不超过 CONTEXT_TARGET_MAX 时直接返回；
    5. 超过时在句末截断，保证不出现半句。
    """
    spans = _sentence_spans(text)

    if not spans:
        # 没有标点没法按句切，退而求其次取字符窗口
        left = max(0, mention_start - 200)
        right = min(len(text), mention_end + 400)
        return sanitize_display(text[left:right])

    # 找实体所在的句
    center = 0
    for i, (start, end) in enumerate(spans):
        if start <= mention_start < end or mention_start <= start < mention_end:
            center = i
            break

    # 初始窗口：前 1 句 + 实体句 + 后 1 句 = 至少 3 句（边界处至少 2 句）
    left = max(0, center - 1)
    right = min(len(spans) - 1, center + 1)

    def current() -> str:
        return sanitize_display(
            text[spans[left][0]:spans[right][1]]
        )

    context = current()

    # 不足目标长度时，继续按句向两侧扩展
    take_left = True
    while (
        _visible_text_len(context) < CONTEXT_TARGET_MIN
        and (left > 0 or right + 1 < len(spans))
    ):
        if take_left and left > 0:
            left -= 1
        elif (not take_left) and right + 1 < len(spans):
            right += 1
        elif left > 0:
            left -= 1
        elif right + 1 < len(spans):
            right += 1

        take_left = not take_left
        context = current()

    # 如果按句扩展仍不够硬性最低，说明源文献本身极短，直接给全部
    if _visible_text_len(context) < CONTEXT_HARD_MIN:
        full = sanitize_display(text)
        if _visible_text_len(full) > _visible_text_len(context):
            context = full

    # 太长时在句末截断，保证至少 2 句
    if _visible_text_len(context) > CONTEXT_TARGET_MAX:
        context = _clip_at_sentence_boundary(
            context, CONTEXT_TARGET_MAX, min_sentences=2
        )

    return context


def excerpt_around(
    entity: dict[str, Any],
    chunk: dict[str, Any] | None,
    document_chunks: list[dict[str, Any]] | None = None,
) -> str:
    """
    从实体实际绑定的 chunk 生成医生复验 context。

    关键点：
    - 不再从“导出后的 PDF”反向识别实体；
    - entity 本身已经带 document_id + chunk_id，所以这里直接使用准确来源；
    - 当前 chunk 太短时，自动拼接同一文献前后 chunk；
    - context 至少 100 字（只要源文献附近有足够文本）。

    旧实现的主要问题是：
        找到实体后，实体后面遇到“第一个句号/分号”就立刻结束，
    所以很多 context 只有一小句。
    """
    if not chunk:
        fallback = sanitize_display(entity.get("evidence_text"))
        return fallback

    current_text = str(chunk.get("text") or "")

    # 先在“当前实体绑定的 chunk”里找实体，确保不会因为同名词在全文多次出现而错位。
    needles = [
        str(entity.get("raw_surface") or ""),
        str(entity.get("name") or ""),
        str(entity.get("semantic_name") or ""),
    ]

    local_index = -1
    needle = ""

    for candidate in needles:
        candidate = candidate.strip()
        if not candidate:
            continue
        local_index = current_text.find(candidate)
        if local_index >= 0:
            needle = candidate
            break

    # 若实体名因为清洗/标准化变化而找不到，使用 evidence 前缀定位。
    if local_index < 0:
        evidence = str(entity.get("evidence_text") or "")
        for prefix_len in (40, 30, 24, 16, 12):
            candidate = evidence[: min(prefix_len, len(evidence))].strip()
            if not candidate:
                continue
            local_index = current_text.find(candidate)
            if local_index >= 0:
                needle = candidate
                break

    # 准备“当前 chunk + 相邻 chunk”窗口。
    joined_text = current_text
    mention_start = local_index if local_index >= 0 else 0

    if document_chunks:
        current_order = int(chunk.get("_index", 1) or 1) - 1
        left_idx = max(0, current_order - CONTEXT_NEIGHBOR_RADIUS)
        right_idx = min(
            len(document_chunks),
            current_order + CONTEXT_NEIGHBOR_RADIUS + 1,
        )

        selected = document_chunks[left_idx:right_idx]
        pieces: list[str] = []
        current_prefix_len = 0

        for absolute_idx, item in enumerate(selected, start=left_idx):
            piece = str(item.get("text") or "")
            if absolute_idx == current_order:
                current_prefix_len = sum(len(x) + 2 for x in pieces)
            pieces.append(piece)

        joined_text = "\n\n".join(pieces)

        if local_index >= 0:
            mention_start = current_prefix_len + local_index
        else:
            # 实体没精确命中时，以当前 chunk 中央为保守中心。
            mention_start = current_prefix_len + max(0, len(current_text) // 2)

    mention_end = mention_start + max(len(needle), 1)

    context = _expand_semantic_window(
        joined_text,
        mention_start=mention_start,
        mention_end=mention_end,
    )

    # 最后一道硬校验：如果仍不足规定字数，就扩大到相邻 chunk 窗口的最大允许长度。
    if _visible_text_len(context) < CONTEXT_HARD_MIN:
        larger = _clip_at_sentence_boundary(
            joined_text, CONTEXT_TARGET_MAX, min_sentences=2
        )
        if _visible_text_len(larger) > _visible_text_len(context):
            context = larger

    # 确保最终上下文不超过目标上限，且在句末截断（至少 2 句）
    if _visible_text_len(context) > CONTEXT_TARGET_MAX:
        context = _clip_at_sentence_boundary(
            context, CONTEXT_TARGET_MAX, min_sentences=2
        )

    return context


def type_display(code: Any) -> str:
    value = sanitize_display(code)
    return f"{TYPE_LABELS.get(value, '未映射类型')}（{value or '空'}）"


def confidence_display(entity: dict[str, Any]) -> str:
    for key in (
        "decision_confidence",
        "calibrated_probability",
        "fusion_probability",
        "candidate_confidence",
    ):
        value = entity.get(key)
        if isinstance(value, (int, float)):
            return f"{key}={value:.3f}"
    return "源记录未提供可用置信度"


def build_review_rows() -> tuple[
    list[dict[str, Any]],
    list[dict[str, Any]],
    dict[str, dict[str, Any]],
]:
    documents, chunks, document_order, chunk_sequences = load_chunks()
    reviewed = load_reviewed_entity_ids()
    source_entities = load_source_entities()

    stats = {
        doc_id: {
            "document_id": doc_id,
            "title": documents[doc_id]["title"],
            "machine_review": 0,
            "human_operated": 0,
            "untouched": 0,
        }
        for doc_id in document_order
    }
    untouched_rows: list[dict[str, Any]] = []

    for entity in source_entities:
        doc_id = entity_doc_id(entity)
        if doc_id not in stats:
            continue
        if str(entity.get("status")) != "review":
            continue
        stats[doc_id]["machine_review"] += 1
        entity_id = str(entity.get("entity_id") or "")
        if entity_id in reviewed:
            stats[doc_id]["human_operated"] += 1
            continue
        stats[doc_id]["untouched"] += 1

        raw_chunk_id = source_chunk_id(entity)
        chunk = chunks.get((doc_id, raw_chunk_id))
        page_start = chunk.get("page_start") if chunk else None
        page_end = chunk.get("page_end") if chunk else None
        if page_start is None:
            pages = "页码缺失"
        elif page_end is None or page_end == page_start:
            pages = str(page_start)
        else:
            pages = f"{page_start}-{page_end}"

        candidate_types: list[str] = []
        for key in (
            "candidate_entity_type",
            "teacher_candidate_type",
            "proposed_entity_type",
            "final_entity_type",
        ):
            value = str(entity.get(key) or "")
            if value and value not in candidate_types:
                candidate_types.append(value)

        effective_type = str(
            entity.get("entity_type")
            or entity.get("proposed_entity_type")
            or entity.get("fusion_entity_type")
            or entity.get("teacher_candidate_type")
            or ""
        )

        mention_start = 10**9
        mention = entity.get("raw_mention_span")
        if isinstance(mention, dict) and isinstance(mention.get("start"), int):
            mention_start = int(mention["start"])

        untouched_rows.append(
            {
                "document_id": doc_id,
                "document_title": documents[doc_id]["title"],
                "document_order": document_order.index(doc_id),
                "entity_id": entity_id,
                "name": sanitize_display(entity.get("name"), limit=160),
                "raw_surface": sanitize_display(entity.get("raw_surface"), limit=160),
                "entity_type": effective_type,
                "entity_type_display": type_display(effective_type),
                "candidate_type_trace": " → ".join(
                    type_display(value) for value in candidate_types
                )
                or "无",
                "chunk_id": f"{doc_id}_{raw_chunk_id}",
                "chunk_order": int(chunk.get("_index", 10**6)) if chunk else 10**6,
                "mention_start": mention_start,
                "pages": pages,
                "section": sanitize_display(
                    entity.get("section_title")
                    or " > ".join(entity.get("section_path") or []),
                    limit=100,
                )
                or "未标注章节",
                "evidence": sanitize_display(
                    entity.get("evidence_text")
                    or entity.get("source_evidence_span", {}).get("text"),
                    limit=420,
                )
                or "源记录未提供证据短句",
                "context": excerpt_around(
                    entity,
                    chunk,
                    chunk_sequences.get(doc_id, []),
                ),
                "confidence": confidence_display(entity),
            }
        )

    untouched_rows.sort(
        key=lambda row: (
            row["document_order"],
            row["chunk_order"],
            row["mention_start"],
            row["entity_id"],
        )
    )

    # 导出前硬校验：复验上下文不得少于硬性最少字数。
    # 若真的存在极短源文本，直接明确报错，而不是生成一份上下文过短的PDF。
    too_short = [
        row
        for row in untouched_rows
        if _visible_text_len(row.get("context", "")) < CONTEXT_HARD_MIN
    ]
    if too_short:
        preview = "\n".join(
            f"- {row['document_title']} | {row['name']} | "
            f"context={_visible_text_len(row.get('context', ''))}字 | "
            f"chunk={row['chunk_id']}"
            for row in too_short[:20]
        )
        raise RuntimeError(
            "存在复验上下文不足 "
            f"{CONTEXT_HARD_MIN} 字的实体，共 {len(too_short)} 条。\n"
            f"{preview}"
        )
    summary_rows = [
        stats[doc_id]
        for doc_id in document_order
    ]
    if (
        EXPECTED_UNTOUCHED is not None
        and len(untouched_rows) != EXPECTED_UNTOUCHED
    ):
        raise RuntimeError(
            f"Expected {EXPECTED_UNTOUCHED} untouched review entities, "
            f"found {len(untouched_rows)}"
        )
    return untouched_rows, summary_rows, documents


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:cantSplit"))
    if node is None:
        node = OxmlElement("w:cantSplit")
        tr_pr.append(node)


def add_valid_page_field(paragraph) -> None:
    prefix = paragraph.add_run("第 ")
    set_run_font(prefix, size=9, color=MUTED)

    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    paragraph._p.append(begin_run)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    instruction_run.append(instruction)
    paragraph._p.append(instruction_run)

    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    paragraph._p.append(separate_run)

    value_run = paragraph.add_run("1")
    set_run_font(value_run, size=9, color=MUTED)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.append(end_run)

    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def add_title_block(
    doc: Document,
    *,
    untouched_count: int,
    document_count: int,
) -> None:
    add_text_paragraph(
        doc,
        "医师复验清单",
        size=24,
        bold=True,
        color=INK,
        after=4,
    )
    count_paragraph = doc.add_paragraph()
    set_paragraph(count_paragraph, before=0, after=3, line_spacing=1.18)
    label_run = count_paragraph.add_run("待复验数量：")
    set_run_font(label_run, size=10, bold=True, color=INK)
    value_run = count_paragraph.add_run(f"{untouched_count}条")
    set_run_font(value_run, size=10, color="333333")


def add_summary_table(doc: Document, summary_rows: list[dict[str, Any]]) -> None:
    heading = doc.add_paragraph(style="Heading 1")
    heading_run = heading.add_run("复验范围汇总")
    set_run_font(heading_run, size=16, bold=True, color=BLUE)
    table = doc.add_table(rows=1, cols=4)
    widths = [3760, 1800, 1800, 2000]
    set_table_geometry(table, widths)
    set_table_borders(table)
    headers = ["文献", "机器需复验", "已人工操作", "本次未操作"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "BFBFBF")
        format_cell_text(cell, header, bold=True, color=INK, size=9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark_repeat_header(table.rows[0])

    for index, item in enumerate(summary_rows, start=1):
        cells = table.add_row().cells
        values = [
            wrap_title(sanitize_display(item["title"], limit=90)),
            str(item["machine_review"]),
            str(item["human_operated"]),
            str(item["untouched"]),
        ]
        for cell_index, value in enumerate(values):
            format_cell_text(cells[cell_index], value, size=9)
            if cell_index != 0:
                cells[cell_index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item["untouched"] > 0:
            set_cell_shading(cells[3], CAUTION_FILL)
        set_row_cant_split(table.rows[-1])


def _clear_cell(cell) -> None:
    """清空单元格，同时保留一个可用段落。"""
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()


def _add_cell_paragraph(
    cell,
    text: str = "",
    *,
    size: float = 10.5,
    bold: bool = False,
    color: str = "222222",
    before: float = 0,
    after: float = 4,
    line_spacing: float = 1.25,
    alignment=None,
):
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    set_paragraph(
        paragraph,
        before=before,
        after=after,
        line_spacing=line_spacing,
        alignment=alignment,
    )
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def _type_name_only(code: Any) -> str:
    value = sanitize_display(code)
    if not value:
        return "未标注"
    return TYPE_LABELS.get(value, value)


def _add_highlighted_context(cell, context: str, entity_name: str, raw_surface: str) -> None:
    """在指南上下文中高亮实体名称；找不到时正常输出全文。"""
    _clear_cell(cell)
    paragraph = cell.paragraphs[0]
    set_paragraph(paragraph, before=0, after=0, line_spacing=1.35)

    needle = ""
    for candidate in (raw_surface, entity_name):
        candidate = sanitize_display(candidate)
        if candidate and candidate in context:
            needle = candidate
            break

    if not needle:
        run = paragraph.add_run(context)
        set_run_font(run, size=10.5, color="222222")
        return

    cursor = 0
    while True:
        index = context.find(needle, cursor)
        if index < 0:
            tail = paragraph.add_run(context[cursor:])
            set_run_font(tail, size=10.5, color="222222")
            break
        if index > cursor:
            normal = paragraph.add_run(context[cursor:index])
            set_run_font(normal, size=10.5, color="222222")
        highlighted = paragraph.add_run(context[index:index + len(needle)])
        set_run_font(highlighted, size=10.5, color="222222")
        highlighted.font.highlight_color = WD_COLOR_INDEX.YELLOW
        cursor = index + len(needle)


def add_entity_overview_table(doc: Document, row: dict[str, Any]) -> None:
    """生成示例中的“实体名称 / 实体类型 / 指南上下文”区域。"""
    table = doc.add_table(rows=2, cols=4)
    widths = [1300, 3600, 1300, 3160]
    set_table_geometry(table, widths)
    set_table_borders(table)

    # 第一行：实体名称、实体类型
    labels = ((0, "实体名称"), (2, "实体类型"))
    for column, label in labels:
        cell = table.cell(0, column)
        set_cell_shading(cell, LIGHT_GRAY)
        format_cell_text(cell, label, bold=True, color=INK, size=10)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    format_cell_text(table.cell(0, 1), row["name"], size=10)
    format_cell_text(
        table.cell(0, 3),
        _type_name_only(row["entity_type"]),
        size=10,
    )

    # 第二行：指南上下文跨 3 列
    label_cell = table.cell(1, 0)
    set_cell_shading(label_cell, LIGHT_GRAY)
    _clear_cell(label_cell)
    p1 = label_cell.paragraphs[0]
    set_paragraph(p1, before=0, after=5, line_spacing=1.2)
    r1 = p1.add_run("指南")
    set_run_font(r1, size=10, bold=True, color=INK)
    p2 = label_cell.add_paragraph()
    set_paragraph(p2, before=0, after=0, line_spacing=1.2)
    r2 = p2.add_run("上下文")
    set_run_font(r2, size=10, bold=True, color=INK)

    context_cell = table.cell(1, 1).merge(table.cell(1, 3))
    set_cell_shading(context_cell, CALLOUT)
    _add_highlighted_context(
        context_cell,
        row["context"],
        row["name"],
        row.get("raw_surface", ""),
    )

    for table_row in table.rows:
        set_row_cant_split(table_row)


def _add_checkbox_line(cell, label: str, *, after: float = 6) -> None:
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    set_paragraph(paragraph, before=0, after=after, line_spacing=1.2)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=10, bold=True, color=INK)
    box_run = paragraph.add_run("  □")
    set_run_font(box_run, size=10, color="222222")


def add_entity_review_form(doc: Document, row: dict[str, Any]) -> None:
    """生成与《医师复验示例》一致的名称、类型和修订说明表单。"""
    table = doc.add_table(rows=3, cols=2)
    widths = [4680, 4680]
    set_table_geometry(table, widths)
    set_table_borders(table)

    # 名称复验：左侧显示机器名称，右侧显示复验操作。
    left_name = table.cell(0, 0)
    _clear_cell(left_name)
    _add_cell_paragraph(left_name, "实体名称：", size=10, bold=True, color=INK, after=9)
    _add_cell_paragraph(left_name, row["name"], size=10.5, after=2, line_spacing=1.35)

    right_name = table.cell(0, 1)
    _clear_cell(right_name)
    _add_checkbox_line(right_name, "接受实体名称", after=7)
    _add_checkbox_line(right_name, "删除实体名称", after=7)
    paragraph = right_name.add_paragraph()
    set_paragraph(paragraph, before=0, after=4, line_spacing=1.2)
    label_run = paragraph.add_run("修改实体名称：")
    set_run_font(label_run, size=10, bold=True, color=INK)
    line_run = paragraph.add_run("____________________")
    set_run_font(line_run, size=10, color="222222")

    # 类型复验：左侧显示机器类型，右侧显示可选类型。
    left_type = table.cell(1, 0)
    _clear_cell(left_type)
    _add_cell_paragraph(left_type, "实体类型", size=10, bold=True, color=INK, after=11)
    _add_cell_paragraph(
        left_type,
        _type_name_only(row["entity_type"]),
        size=10.5,
        after=2,
    )

    right_type = table.cell(1, 1)
    _clear_cell(right_type)
    _add_checkbox_line(right_type, "接受实体类型：", after=8)
    _add_cell_paragraph(right_type, "修改实体类型：", size=10, bold=True, color=INK, after=6)

    choices = [
        ("疾病名", "确诊名"),
        ("病因", "发病机制"),
        ("临床表现", "检查指标"),
        ("治疗原则", "治疗方案"),
    ]
    for left_choice, right_choice in choices:
        paragraph = right_type.add_paragraph()
        set_paragraph(paragraph, before=0, after=5, line_spacing=1.2)
        left_run = paragraph.add_run(f"{left_choice}  □")
        set_run_font(left_run, size=10, color="222222")
        spacer = " " * max(6, 18 - len(left_choice) * 2)
        middle_run = paragraph.add_run(spacer)
        set_run_font(middle_run, size=10, color="222222")
        right_run = paragraph.add_run(f"{right_choice}  □")
        set_run_font(right_run, size=10, color="222222")

    # 修订说明跨两列，并保留书写空间。
    notes_cell = table.cell(2, 0).merge(table.cell(2, 1))
    _clear_cell(notes_cell)
    _add_cell_paragraph(notes_cell, "修订说明：", size=10, bold=True, color=INK, after=16)
    for _ in range(3):
        _add_cell_paragraph(notes_cell, "", size=10, after=12)

    for table_row in table.rows:
        set_row_cant_split(table_row)


def add_entity_card(
    doc: Document,
    row: dict[str, Any],
    number: int,
    *,
    document_title: str,
    document_entity_count: int,
) -> None:
    title = doc.add_paragraph(style="Heading 1")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title_run = title.add_run(f"{wrap_title(document_title)}（{document_entity_count}条）")
    set_run_font(title_run, size=16, bold=True, color=INK)

    heading = doc.add_paragraph()
    set_paragraph(heading, before=0, after=11, line_spacing=1.25, keep_with_next=True)
    number_run = heading.add_run(f"实体 {number:02d} · ")
    set_run_font(number_run, size=12, bold=True, color=INK)

    add_entity_overview_table(doc, row)
    add_text_paragraph(doc, "", size=2, after=4)
    add_entity_review_form(doc, row)

def add_section_signature(doc: Document, title: str, count: int) -> None:
    sig_heading = doc.add_paragraph(style="Heading 3")
    sig_heading_run = sig_heading.add_run("文献复验签署")
    set_run_font(sig_heading_run, size=12, bold=True, color=DARK_BLUE)
    table = doc.add_table(rows=2, cols=4)
    widths = [1400, 3280, 1400, 3280]
    values = [
        ("判定医师", "签名：", "科室 / 职称", ""),
        ("复验数量", f"{count}条", "判定日期", "____年__月__日"),
    ]
    for row_index, row_values in enumerate(values):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            if column_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                format_cell_text(cell, value, bold=True, color=INK, size=9.0)
            else:
                format_cell_text(cell, value, size=9.0)
        set_row_cant_split(table.rows[row_index])
    set_table_geometry(table, widths)
    set_table_borders(table)


def configure_document_header(doc: Document) -> None:
    configure_page(doc)
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    set_paragraph(
        paragraph,
        before=0,
        after=0,
        line_spacing=1.0,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    run = paragraph.add_run("CmePlatform  |  医师复验清单")
    set_run_font(run, size=9, bold=True, color=MUTED)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    set_paragraph(
        paragraph,
        before=0,
        after=0,
        line_spacing=1.0,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    add_valid_page_field(paragraph)


def build_document() -> Path:
    untouched_rows, summary_rows, _ = build_review_rows()
    by_document: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in untouched_rows:
        by_document[row["document_id"]].append(row)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_document_header(doc)

    props = doc.core_properties
    props.title = "医师复验清单：机器待复验且人工未操作实体"
    props.subject = "当前批次机器待复验且当前人工结论为 pending 的实体"
    props.author = "CmePlatform"
    props.keywords = "知识复验, 医师判定, 未操作实体"

    add_title_block(
        doc,
        untouched_count=len(untouched_rows),
        document_count=len(summary_rows),
    )
    add_summary_table(doc, summary_rows)

    number = 1
    for summary in summary_rows:
        doc_id = summary["document_id"]
        rows = by_document.get(doc_id, [])
        if not rows:
            continue
        for row_index, row in enumerate(rows):
            add_entity_card(
                doc,
                row,
                number,
                document_title=summary["title"],
                document_entity_count=len(rows),
            )
            number += 1

    rendered_count = number - 1
    if rendered_count != len(untouched_rows):
        raise RuntimeError(
            "Rendered entity count mismatch: "
            f"rendered {rendered_count}, expected {len(untouched_rows)}"
        )

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
